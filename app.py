import streamlit as st
import datetime
import json
import time
from io import BytesIO
from sentence_transformers import SentenceTransformer
import weaviate
from weaviate.auth import AuthApiKey
import groq
from docx import Document
import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- CONFIGURATION ---

# Load sensitive data from secrets
WEAVIATE_URL = st.secrets["api_keys"]["WEAVIATE_URL"]
WEAVIATE_API_KEY = st.secrets["api_keys"]["WEAVIATE_API_KEY"]
GROQ_API_KEY = st.secrets["api_keys"]["GROQ_API_KEY"]
INDEX_NAME = "tcmapp"  # Adjust as needed if your Weaviate index name is different

GOOGLE_SHEETS_CREDENTIALS = st.secrets["gcp_service_account"]
SHEET_ID = st.secrets["google_sheets"]["sheet_id"]

# Initialize Weaviate client
auth = AuthApiKey(api_key=WEAVIATE_API_KEY)
weaviate_client = weaviate.Client(WEAVIATE_URL, auth_client_secret=auth)

# --- UTILITY FUNCTIONS ---

def calculate_age(born):
    today = datetime.date.today()
    return today.year - born.year - ((today.month, today.day) < (born.month, born.day))

def clear_patient_data():
    st.session_state.patient_info = {}
    st.session_state.generated_report = None
    st.success("Patient data has been cleared.")

# Google API Initialization
def initialize_google_services():
    creds = Credentials.from_service_account_info(
        GOOGLE_SHEETS_CREDENTIALS,
        scopes=[
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
    )
    sheets_service = build("sheets", "v4", credentials=creds)
    docs_service = build("docs", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds)
    return sheets_service, docs_service, drive_service

# Patient Search Function
def search_patient(sheets_service, name):
    try:
        result = (
            sheets_service.spreadsheets()
            .values()
            .get(spreadsheetId=SHEET_ID, range="A:Z")
            .execute()
        )
        values = result.get("values", [])
        headers = values[0]
        for row in values[1:]:
            if row[0].lower() == name.lower():
                patient_data = dict(zip(headers, row))
                # Check if there's a report ID and fetch the report content
                if 'Report ID' in patient_data and patient_data['Report ID']:
                    report_content = get_report_content(docs_service, patient_data['Report ID'])
                    patient_data['Report Content'] = report_content
                return patient_data
        return None
    except Exception as e:
        st.error(f"Error searching for patient: {e}")
        return None

# Patient Save/Update Function
def save_or_update_patient(sheets_service, patient_data):
    try:
        result = sheets_service.spreadsheets().values().get(spreadsheetId=SHEET_ID, range="A:A").execute()
        names = result.get('values', [])
        row_index = next((i for i, name in enumerate(names) if name and name[0] == patient_data["Patient Name"]), None)

        values = [list(patient_data.values())]
        if row_index is not None:
            range_name = f'Sheet1!A{row_index + 1}'
            sheets_service.spreadsheets().values().update(
                spreadsheetId=SHEET_ID,
                range=range_name,
                valueInputOption="USER_ENTERED",
                body={"values": values}
            ).execute()
        else:
            sheets_service.spreadsheets().values().append(
                spreadsheetId=SHEET_ID,
                range="Sheet1!A1",
                valueInputOption="USER_ENTERED",
                body={"values": values}
            ).execute()
        st.success("Patient data saved/updated successfully")
    except Exception as e:
        st.error(f"Error saving/updating patient data: {e}")

# Function to get all patients
def get_all_patients(sheets_service):
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=SHEET_ID,
            range="A2:A"  # Assuming patient names are in column A, starting from row 2
        ).execute()
        values = result.get('values', [])
        return [name[0] for name in values if name]  # Flatten the list and remove empty entries
    except Exception as e:
        st.error(f"Error fetching patient list: {e}")
        return []

# Function to save report to Google Docs
def save_report_to_docs(docs_service, drive_service, patient_name, report_content):
    try:
        # Create a folder for TCM reports if it doesn't exist
        folder_name = 'TCM Diagnostic Reports'
        folder_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        # Check if folder already exists
        query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
        results = drive_service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
        folders = results.get('files', [])
        
        if not folders:
            # Create the folder if it doesn't exist
            folder = drive_service.files().create(body=folder_metadata, fields='id').execute()
            folder_id = folder.get('id')
        else:
            # Use existing folder
            folder_id = folders[0]['id']

        # Create a new Google Doc
        doc = {
            'name': f"TCM Diagnostic Report - {patient_name}",
            'parents': [folder_id]
        }
        doc = drive_service.files().create(body=doc, fields='id, webViewLink').execute()
        doc_id = doc.get('id')
        doc_link = doc.get('webViewLink')

        # Write content to the document
        requests = [
            {
                'insertText': {
                    'location': {
                        'index': 1,
                    },
                    'text': report_content
                }
            }
        ]

        # Debug: Print the size of the content
        st.write(f"Debug: Report content size: {len(report_content)} characters")

        # Try to update the document content
        try:
            docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': requests}).execute()
        except Exception as update_error:
            st.error(f"Error updating document content: {update_error}")
            # If update fails, try to insert content in smaller chunks
            chunk_size = 1000000  # Google Docs has a limit of about 1MB per request
            for i in range(0, len(report_content), chunk_size):
                chunk = report_content[i:i+chunk_size]
                chunk_request = [
                    {
                        'insertText': {
                            'location': {
                                'index': i + 1,
                            },
                            'text': chunk
                        }
                    }
                ]
                docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': chunk_request}).execute()

        # Set permissions to anyone with the link can view
        drive_service.permissions().create(
            fileId=doc_id,
            body={'type': 'anyone', 'role': 'reader'},
            fields='id'
        ).execute()

        return doc_id, doc_link
    except Exception as e:
        st.error(f"Error saving report to Google Docs: {str(e)}")
        if hasattr(e, 'content'):
            st.error(f"Error details: {e.content}")
        return None, None

# Function to get report content from Google Docs
def get_report_content(docs_service, doc_id):
    try:
        document = docs_service.documents().get(documentId=doc_id).execute()
        content = document.get('body').get('content')
        full_text = ""
        for elem in content:
            if 'paragraph' in elem:
                for para_elem in elem['paragraph']['elements']:
                    if 'textRun' in para_elem:
                        full_text += para_elem['textRun']['content']
        return full_text
    except Exception as e:
        st.error(f"Error fetching report from Google Docs: {e}")
        return None

# --- STYLING ---

st.markdown(
    """
    <style>
    /* Custom styles for Streamlit app */
    body {
        font-family: 'Arial', sans-serif;
        color: #f6f6f6;
    }

    h1 {
        font-size: 2.5em;
        text-align: center;
        color: #2962ff;
    }

    .stButton>button {
        background-color: #2962ff;
        color: white;
        border: none;
        padding: 10px 20px;
        border-radius: 5px;
        cursor: pointer;
    }

    .stButton>button:hover {
        background-color: #1e4db7;
    }

    .stTextInput>div>div>input {
        background-color: #0c0e15;
        color: #f6f6f6;
        border: 1px solid #2962ff;
        padding: 10px;
        border-radius: 5px;
    }

    .stTextInput>div>div>input:focus {
        border-color: #1e4db7;
        outline: none;
    }

    .stDataFrame {
        background-color: #0c0e15;
        color: #f6f6f6;
        border: 1px solid #2962ff;
        border-radius: 5px;
        padding: 10px;
    }

    .stExpander {
        background-color: #0c0e15;
        border: 1px solid #2962ff;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- QUERY & REPORT GENERATION FUNCTIONS ---

def query_weaviate(query_text, top_k=5):
    if weaviate_client is None:
        raise ValueError("Weaviate client is not initialized")
    query_vector = embedding_model.encode(query_text).tolist()
    near_vector = {"vector": query_vector}
    results = weaviate_client.query.get("TCMApp", ["text"]).with_near_vector(near_vector).with_limit(top_k).do()

    if 'data' in results and 'Get' in results['data'] and 'TCMApp' in results['data']['Get']:
        return results['data']['Get']['TCMApp']
    else:
        st.error(f"Unexpected response format: {results}")
        return []

def generate_diagnostic_report_part(system_message, user_message, context):
    try:
        response = groq_client.chat.completions.create(
            model="mixtral-8x7b-32768", 
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": f"Context: {context}\n\nUser Query: {user_message}"}
            ],
            temperature=0.7,
            max_tokens=4000,
            top_p=1,
            stop=None,
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        return None

def generate_diagnostic_report(context, user_input):
    patient_info = json.loads(user_input)
    patient_name = patient_info.get('Patient Name', 'Patient')
    patient_age = patient_info.get('Age', 'Unknown age')

    system_message = f"""You are a TCM practitioner tasked with generating a diagnostic report. Your knowledge is strictly limited to the information provided in the context. Do not use any external knowledge or make assumptions beyond what is explicitly stated in the context or patient information.

    You are generating a report for {patient_name}, a {patient_age}-year-old patient. Ensure your report is based solely on the provided context and patient information. If you cannot find relevant information in the context to address a particular aspect of the report, state that there is insufficient information to make a determination on that point."""

    report_sections = [
        "1. Patient Overview",
        "2. TCM Diagnosis",
        "3. Pattern Differentiation",
        "4. Treatment Principle and Plan",
        "5. Acupuncture Point Prescription",
        "6. Herbal Formula Recommendation",
        "7. Lifestyle and Dietary Advice",
        "8. Prognosis and Follow-up Recommendations"
    ]

    document = Document()
    document.add_heading(f"TCM Diagnostic Report for {patient_name}", 0)

    progress_bar = st.progress(0)

    for i, section in enumerate(report_sections):
        user_message = f"""
        Based strictly on the following patient information and context, generate the {section} of the TCM diagnostic report.
        Only use information explicitly provided in the context or patient information. If there's insufficient information for any part of this section, clearly state this limitation.
        Patient Information: {user_input}
        """

        query_results = query_weaviate(user_message)
        section_context = "\n".join([result['text'] for result in query_results])

        with st.spinner(f"Generating {section}..."):
            section_content = generate_diagnostic_report_part(system_message, user_message, section_context)
            if section_content:
                document.add_heading(section, level=1)
                document.add_paragraph(section_content)
            else:
                st.warning(f"Failed to generate {section}. Moving to the next section.")

        progress_bar.progress((i + 1) / len(report_sections))
        time.sleep(1)  # Small delay for better UX

    return document

# --- MAIN APP ---

# Initialize resources
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
groq_client = groq.Client(api_key=GROQ_API_KEY)
sheets_service, docs_service, drive_service = initialize_google_services()

# Session state for better UX
if "patient_info" not in st.session_state:
    st.session_state.patient_info = {}
if "generated_report" not in st.session_state:
    st.session_state.generated_report = None

def patient_info_page():
    st.title("AcuAssist: Your AI-Powered TCM Diagnostic Assistant")

    # Patient Selection Dropdown
    all_patients = get_all_patients(sheets_service)
    selected_patient = st.selectbox("Select Patient", ["New Patient"] + all_patients)

    if selected_patient != "New Patient":
        # Load the selected patient's data
        patient_data = search_patient(sheets_service, selected_patient)
        if patient_data:
            st.session_state.patient_info = patient_data
            st.success(f"Loaded data for patient: {selected_patient}")
            
            # Display the report link if it exists
            if 'Report ID' in patient_data and 'Report Link' in patient_data:
                st.subheader("Existing TCM Diagnostic Report")
                st.markdown(f"[View Report]({patient_data['Report Link']})")
        else:
            st.error(f"Failed to load data for patient: {selected_patient}")
    else:
        # Clear the form for a new patient
        clear_patient_data()

    # Patient Search
    with st.expander("Patient Search", expanded=False):
        search_col1, search_col2 = st.columns([3, 1])
        with search_col1:
            search_name = st.text_input("Search Patient by Name")
        with search_col2:
            search_button = st.button("Search")

        if search_button:
            found_patient = search_patient(sheets_service, search_name)
            if found_patient:
                st.success(f"Patient '{search_name}' found!")
                st.session_state.patient_info = found_patient
                st.experimental_rerun()  # Refresh the page with the found data
            else:
                st.warning(f"No patient found with name '{search_name}'")

    # Basic Information
    with st.expander("Basic Information", expanded=True):
        patient_data = st.session_state.get('patient_info', {})

        col1, col2, col3 = st.columns(3)
        with col1:
            name = st.text_input("Patient Name", key="name", value=patient_data.get('Patient Name', ''))
        with col2:
            gender = st.selectbox("Gender", ["Male", "Female", "Other"], key="gender", index=["Male", "Female", "Other"].index(patient_data.get('Gender', 'Male')))
        with col3:
            dob = patient_data.get('Date of Birth (DD/MM/YY)', '')
            if dob:
                try:
                    dob = datetime.datetime.strptime(dob, "%d/%m/%y")
                except ValueError:
                    st.error("Invalid date format in stored data. Using default value.")
                    dob = datetime.datetime.now()
            else:
                dob = datetime.datetime.now()
            
            dob = st.date_input("Date of Birth", value=dob)
            dob_str = dob.strftime("%d/%m/%y")
            age = calculate_age(dob)
            st.write(f"Patient Age: {age} years")

    # Chief Complaint
    with st.expander("Chief Complaint", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            chief_complaint = st.text_area("Main Complaint", key="chief_complaint", value=patient_data.get('Chief Complaint', ''))
        with col2:
            complaint_duration = st.text_input("Duration of Complaint", key="complaint_duration", value=patient_data.get('Duration of Complaint', ''))

    # TCM Four Diagnostic Methods
    with st.expander("TCM Four Diagnostic Methods", expanded=True):
        st.subheader("1. Inspection (望 wàng)")
        col1, col2, col3 = st.columns(3)
        with col1:
            complexion = st.text_input("Complexion", key="complexion", value=patient_data.get('Complexion', ''))
        with col2:
            tongue_color_options = ["Not observed", "Pale", "Red", "Dark Red", "Purple", "Bluish Purple"]
            stored_tongue_color = patient_data.get('Tongue Color', 'Not observed')
            tongue_color_index = tongue_color_options.index(stored_tongue_color) if stored_tongue_color in tongue_color_options else 0
            tongue_color = st.selectbox("Tongue Color", tongue_color_options, key="tongue_color", index=tongue_color_index)
        with col3:
            tongue_coating_options = ["Not observed", "Thin White", "Thick White", "Yellow", "Grey", "Black"]
            stored_tongue_coating = patient_data.get('Tongue Coating', 'Not observed')
            tongue_coating_index = tongue_coating_options.index(stored_tongue_coating) if stored_tongue_coating in tongue_coating_options else 0
            tongue_coating = st.selectbox("Tongue Coating", tongue_coating_options, key="tongue_coating", index=tongue_coating_index)
        tongue_shape = st.text_input("Tongue Shape and Features", key="tongue_shape", value=patient_data.get('Tongue Shape and Features', ''))

        st.subheader("2. Auscultation and Olfaction (聞 wén)")
        col1, col2 = st.columns(2)
        with col1:
            voice_sound = st.text_input("Voice Sound", key="voice_sound", value=patient_data.get('Voice Sound', ''))
        with col2:
            breath_odor = st.text_input("Breath Odor", key="breath_odor", value=patient_data.get('Breath Odor', ''))

        st.subheader("3. Inquiry (問 wèn)")
        col1, col2 = st.columns(2)
        with col1:
            cold_heat_options = ["Aversion to Cold", "Aversion to Heat", "Alternating Cold and Heat", "Normal"]
            stored_cold_heat = patient_data.get('Cold/Heat Sensation', 'Normal')
            cold_heat_index = cold_heat_options.index(stored_cold_heat) if stored_cold_heat in cold_heat_options else 3  # Default to 'Normal'
            cold_heat_sensation = st.selectbox("Cold/Heat Sensation", cold_heat_options, key="cold_heat_sensation", index=cold_heat_index)
            
            sweating = st.text_input("Sweating", key="sweating", value=patient_data.get('Sweating', ''))
            appetite = st.text_input("Appetite and Thirst", key="appetite", value=patient_data.get('Appetite and Thirst', ''))
        with col2:
            sleep = st.text_input("Sleep Pattern", key="sleep", value=patient_data.get('Sleep Pattern', ''))
            bowel_movements = st.text_input("Bowel Movements", key="bowel_movements", value=patient_data.get('Bowel Movements', ''))
            urination = st.text_input("Urination", key="urination", value=patient_data.get('Urination', ''))
        pain = st.text_area("Pain (location, nature, factors that alleviate or aggravate)", key="pain", value=patient_data.get('Pain (location, nature, factors that alleviate or aggravate)', ''))

        st.subheader("4. Palpation (切 qiè)")
        col1, col2 = st.columns(2)
        with col1:
            stored_pulse_rate = patient_data.get('Pulse Rate (BPM)', '70')
            try:
                default_pulse_rate = int(float(stored_pulse_rate))
                default_pulse_rate = max(40, min(200, default_pulse_rate))  # Ensure it's within the allowed range
            except ValueError:
                default_pulse_rate = 70  # Default to 70 if conversion fails
            
            pulse_rate = st.number_input("Pulse Rate (BPM)", key="pulse_rate", min_value=40, max_value=200, value=default_pulse_rate)
        
        with col2:
            pulse_quality_options = ["Floating", "Sinking", "Slow", "Rapid", "Strong", "Weak", "Wiry", "Slippery", "Rough"]
            stored_pulse_quality = patient_data.get('Pulse Quality', '')

            if isinstance(stored_pulse_quality, str):
                default_pulse_quality = [item.strip() for item in stored_pulse_quality.split(',') if item.strip() in pulse_quality_options]
            elif isinstance(stored_pulse_quality, list):
                default_pulse_quality = [item for item in stored_pulse_quality if item in pulse_quality_options]
            else:
                default_pulse_quality = []

            pulse_quality = st.multiselect("Pulse Quality", pulse_quality_options, key="pulse_quality", default=default_pulse_quality)

    # Additional TCM Diagnostic Information
    with st.expander("Additional TCM Diagnostic Information", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            emotions = st.text_area("Emotional State", key="emotions", value=patient_data.get('Emotional State', ''))
        with col2:
            lifestyle = st.text_area("Lifestyle Factors (diet, exercise, stress, etc.)", key="lifestyle", value=patient_data.get('Lifestyle Factors (diet, exercise, stress, etc.)', ''))
        medical_history = st.text_area("Relevant Medical History", key="medical_history", value=patient_data.get('Relevant Medical History', ''))

    # Update session state with current form values
    st.session_state.patient_info.update({
        'Patient Name': name,
        'Date of Birth (DD/MM/YY)': dob_str,
        'Age': str(age) if age is not None else '',
        'Gender': gender,
        'Chief Complaint': chief_complaint,
        'Duration of Complaint': complaint_duration,
        'Complexion': complexion,
        'Tongue Color': tongue_color,
        'Tongue Coating': tongue_coating,
        'Tongue Shape and Features': tongue_shape,
        'Voice Sound': voice_sound,
        'Breath Odor': breath_odor,
        'Cold/Heat Sensation': cold_heat_sensation,
        'Sweating': sweating,
        'Appetite and Thirst': appetite,
        'Sleep Pattern': sleep,
        'Bowel Movements': bowel_movements,
        'Urination': urination,
        'Pain (location, nature, factors that alleviate or aggravate)': pain,
        'Pulse Rate (BPM)': str(pulse_rate),
        'Pulse Quality': ', '.join(pulse_quality),  # Comma-separated string for multiselect
        'Emotional State': emotions,
        'Lifestyle Factors (diet, exercise, stress, etc.)': lifestyle,
        'Relevant Medical History': medical_history
    })

def view_report_page():
    st.title("AcuAssist: TCM Diagnostic Report")
    if st.session_state.generated_report:
        doc = st.session_state.generated_report
        for paragraph in doc.paragraphs:
            st.write(paragraph.text)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Download Report as Word Document",
            data=buffer,
            file_name="tcm_diagnostic_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("No report has been generated yet. Please go to the 'Patient Information' page to enter patient data and generate a report.")

def main():
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Patient Information", "View Report"])

    if page == "Patient Information":
        patient_info_page()
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("Save Patient Information"):
                if 'Patient Name' in st.session_state.patient_info and st.session_state.patient_info['Patient Name']:
                    save_or_update_patient(sheets_service, st.session_state.patient_info)
                else:
                    st.error("Please enter patient name before saving")
        with col2:
            if st.button("Generate TCM Diagnostic Report"):
                if len(st.session_state.patient_info) > 10:  # Check if enough info is filled
                    try:
                        serializable_patient_info = st.session_state.patient_info.copy()
                        serializable_patient_info['Age'] = calculate_age(datetime.datetime.strptime(serializable_patient_info['Date of Birth (DD/MM/YY)'], "%d/%m/%y"))

                        user_input = json.dumps(serializable_patient_info, indent=2)

                        # Query Weaviate for initial context
                        query_results = query_weaviate(user_input)
                        context = "\n".join([result['text'] for result in query_results])

                        start_time = time.time()
                        report = generate_diagnostic_report(context, user_input)
                        end_time = time.time()

                        if report:
                            st.success(f"Report generated in {end_time - start_time:.2f} seconds")
                            st.session_state.generated_report = report
                            
                            # Save report to Google Docs
                            report_content = "\n".join([para.text for para in report.paragraphs])
                            doc_id, doc_link = save_report_to_docs(docs_service, drive_service, serializable_patient_info['Patient Name'], report_content)
                            
                            if doc_id and doc_link:
                                st.session_state.patient_info['Report ID'] = doc_id
                                st.session_state.patient_info['Report Link'] = doc_link
                                save_or_update_patient(sheets_service, st.session_state.patient_info)
                                st.write("TCM Diagnostic Report generated and saved successfully.")
                                st.markdown(f"[View Report]({doc_link})")
                            else:
                                st.error("Failed to save the report to Google Docs.")
                        else:
                            st.error("Failed to generate the report. Please try again.")
                    except Exception as e:
                        st.error(f"An error occurred during report generation: {str(e)}")
                else:
                    st.warning("Please fill in more patient information before generating a report.")
        with col3:
            if st.button("Clear Form"):
                clear_patient_data()
                st.experimental_rerun()  # Refresh the page after clearing
    elif page == "View Report":
        view_report_page()

if __name__ == "__main__":
    main()