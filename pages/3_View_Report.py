import streamlit as st
import re
from urllib.parse import quote
import streamlit.components.v1 as components
from io import BytesIO
from docx import Document as DocxDocument
import datetime

st.title("View TCM Diagnostic Report")

# Function to create DOCX report
def create_docx_report(report_text, mermaid_diagrams):
    doc = DocxDocument()
    doc.add_heading('TCM Diagnostic Report', 0)

    # Add report text
    for paragraph in report_text.split('\n'):
        doc.add_paragraph(paragraph)

    # Add Mermaid diagrams
    if mermaid_diagrams:
        doc.add_heading('TCM Pattern Differentiation Diagrams', level=1)
        for i, diagram in enumerate(mermaid_diagrams):
            doc.add_paragraph(f"Diagram {i+1}:")
            doc.add_paragraph(diagram)

    return doc

if 'generated_report' in st.session_state:
    report = st.session_state.generated_report
    
    # Extract Mermaid diagrams
    pattern = r'```mermaid\n(.*?)```'
    matches = re.findall(pattern, report, re.DOTALL)
    
    # Remove the Mermaid code blocks from the text
    report_without_diagrams = re.sub(pattern, '', report, flags=re.DOTALL)
    
    # Display the report text
    st.markdown(report_without_diagrams)
    
    # Render Mermaid diagrams
    st.header("TCM Pattern Differentiation Diagrams")
    for i, match in enumerate(matches):
        # Create a unique key for each Mermaid diagram
        key = f"mermaid_diagram_{i}"
        
        # Encode the Mermaid diagram content
        encoded_diagram = quote(match.strip())
        
        # Render the Mermaid diagram using a custom component
        components.html(
            f"""
            <script src="https://cdn.jsdelivr.net/npm/mermaid@10.9.1/dist/mermaid.min.js"></script>
            <div id="{key}_container">
                <div class="mermaid" id="{key}"></div>
            </div>
            <pre id="{key}_fallback" style="display:none; white-space: pre-wrap; word-wrap: break-word;">
                {match.strip()}
            </pre>
            <script>
                mermaid.initialize({{ startOnLoad: false }});
                function renderDiagram() {{
                    var encodedDiagram = "{encoded_diagram}";
                    var decodedDiagram = decodeURIComponent(encodedDiagram);
                    mermaid.render('{key}_svg', decodedDiagram).then(function(result) {{
                        document.getElementById('{key}').innerHTML = result.svg;
                    }}).catch(function(error) {{
                        console.error('Error rendering Mermaid diagram:', error);
                        document.getElementById('{key}_container').style.display = 'none';
                        document.getElementById('{key}_fallback').style.display = 'block';
                    }});
                }}
                function adjustHeight() {{
                    var diagram = document.getElementById("{key}");
                    var svg = diagram.querySelector('svg');
                    if (svg) {{
                        var height = svg.getBoundingClientRect().height;
                        diagram.style.height = height +diagram.style.height = height + 'px';
                        parent.postMessage({{type: 'setFrameHeight', height: height + 50}}, '*');
                    }}
                }}
                renderDiagram();
                window.addEventListener('load', adjustHeight);
                new MutationObserver(adjustHeight).observe(document.getElementById("{key}"), {{ attributes: true, childList: true, subtree: true }});
            </script>
            """,
            height=600,  # Initial height, will be adjusted by JavaScript
            scrolling=True
        )
        st.write("\n")

    # Create Word document
    doc = create_docx_report(report_without_diagrams, matches)
    
    # Save the document to a BytesIO object
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    # Create a download button for the Word document
    st.download_button(
        label="Download Report as Word Document",
        data=docx_io,
        file_name=f"TCM_Report_{st.session_state.patient_info['name']}_{datetime.date.today().strftime('%Y-%m-%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.warning("No generated report found. Please generate a report in the 'Generate Report' page first.")
